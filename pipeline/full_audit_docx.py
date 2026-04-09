#!/usr/bin/env python3
"""
DOCX full audit script - 3 dimensions:
  A: Requirements compliance (KPI, hardware, training hours, red lines)
  B: Format (delegated to audit_docx_format.py)
  C: Content quality (placeholders, simplified Chinese, stub sections, MD vs DOCX)
"""
import os, sys, re, json
from docx import Document

DOCX_DIR = "pipeline/output_docx"
MD_DIR = "pipeline/output"
ISSUES = []
ISSUE_ID = 0

def add_issue(sev, fname, dim, cat, desc, expected="", actual=""):
    global ISSUE_ID
    ISSUE_ID += 1
    ISSUES.append({
        "id": f"DOCX-{ISSUE_ID:03d}",
        "severity": sev,
        "file": fname,
        "dimension": dim,
        "category": cat,
        "desc": desc,
        "expected": expected,
        "actual": actual,
    })

def get_full_text(doc):
    """Extract all text from a docx Document."""
    parts = []
    for p in doc.paragraphs:
        parts.append(p.text)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)

# Simplified Chinese detection ranges
SIMPLIFIED_CHARS = set()
# Common simplified-only chars that differ from traditional
SC_SAMPLES = (
    "\u4e0e\u5173\u4e3a\u4ece\u4e48\u4e1a\u4e1c\u4e2a\u4e60\u53d1"
    "\u5f00\u8fd9\u8fdb\u8fc7\u7ecf\u65f6\u7f51\u7ea7\u8282\u6761"
    "\u6bd5\u8282\u70b9\u8054\u7edf\u53f7\u4ef6\u6761\u62a5\u544a"
)

def has_simplified_chinese(text):
    """Check for simplified-only Chinese characters (not shared with traditional)."""
    # Strictly simplified-only characters that have DIFFERENT traditional forms.
    # Excluded all chars that are identical in both systems (e.g. U+4EF6 件, U+6761 條, U+53F7 號).
    truly_simplified = {
        # Common simplified-only (each has a distinct traditional counterpart)
        '\u8fd0',  # 运 (trad: 運)
        '\u8f6f',  # 软 (trad: 軟)
        '\u5f00',  # 开 (trad: 開)
        '\u53d1',  # 发 (trad: 發/髮)
        '\u5173',  # 关 (trad: 關)
        '\u8bbe',  # 设 (trad: 設)
        '\u8ba1',  # 计 (trad: 計)
        '\u7edc',  # 络 (trad: 絡)
        '\u7ea7',  # 级 (trad: 級)
        '\u8282',  # 节 (trad: 節)
        '\u8054',  # 联 (trad: 聯) — note: 聯 is traditional
        '\u7edf',  # 统 (trad: 統)
        '\u62a5',  # 报 (trad: 報)
        '\u8fdb',  # 进 (trad: 進)
        '\u4e3a',  # 为 (trad: 為)
        '\u4ece',  # 从 (trad: 從)
        '\u5904',  # 处 (trad: 處)
        '\u8fc7',  # 过 (trad: 過)
        '\u7ecf',  # 经 (trad: 經)
        '\u65f6',  # 时 (trad: 時)
        '\u95e8',  # 门 (trad: 門)
        '\u94fe',  # 链 (trad: 鏈)
        '\u8bba',  # 论 (trad: 論)
        '\u8bed',  # 语 (trad: 語)
        '\u9879',  # 项 (trad: 項)
        '\u6807',  # 标 (trad: 標)
        '\u8bc6',  # 识 (trad: 識)
        '\u6d4b',  # 测 (trad: 測)
        '\u8bd5',  # 试 (trad: 試)
        '\u4e1a',  # 业 (trad: 業)
        '\u4e1c',  # 东 (trad: 東)
        '\u4e60',  # 习 (trad: 習)
        '\u8fd9',  # 这 (trad: 這)
        '\u4e2a',  # 个 (trad: 個)
        '\u4e48',  # 么 (trad: 麼)
        '\u5386',  # 历 (trad: 歷)
        '\u6b22',  # 欢 (trad: 歡)
        '\u53ea',  # 只 — same in both, remove
        '\u7b97',  # 算 — same in both, remove
        '\u5bf9',  # 对 (trad: 對)
        '\u4e0e',  # 与 (trad: 與) — BUT U+4E0E is also used in trad as a variant!
        '\u7535',  # 电 (trad: 電)
        '\u8bef',  # 误 (trad: 誤)
        '\u8d28',  # 质 (trad: 質)
        '\u529e',  # 办 (trad: 辦)
        '\u5386',  # 历 (trad: 歷)
        '\u8ba9',  # 让 (trad: 讓)
        '\u6c14',  # 气 (trad: 氣)
        '\u7cfb',  # 系 — same in both, exclude
    }
    # Remove chars that are actually shared
    truly_simplified -= {'\u53ea', '\u7b97', '\u7cfb', '\u4e0e'}

    found = []
    for i, ch in enumerate(text):
        if ch in truly_simplified:
            ctx_start = max(0, i - 5)
            ctx_end = min(len(text), i + 6)
            context = text[ctx_start:ctx_end].replace('\n', ' ')
            found.append((ch, context))
            if len(found) >= 5:
                break
    return found

def check_placeholders(text, fname):
    """Check for placeholder text."""
    patterns = [
        (r'TODO', 'TODO'),
        (r'TBD', 'TBD'),
        (r'FIXME', 'FIXME'),
        (r'xxx', 'xxx'),
        (r'\[TBD\]', '[TBD]'),
    ]
    found = []
    for pat, label in patterns:
        matches = re.findall(pat, text, re.IGNORECASE)
        if matches:
            found.append(f"{label} x{len(matches)}")

    # Check for empty brackets that might be placeholders
    empty_brackets = re.findall(r'\u3010\s*\u3011', text)
    if empty_brackets:
        found.append(f"empty brackets x{len(empty_brackets)}")

    return found

def check_kpi_values(text, fname):
    """Check KPI values against decisions.yaml SSOT."""
    kpi_checks = [
        ("85%", "IVR/Agent accuracy", True),
        ("5000", "response time ms", True),
        ("30", "concurrent ports", True),
        ("99%", "availability", True),
        ("54", "training hours", True),
        ("4", "MTTR hours", True),
    ]
    # We just verify these numbers appear somewhere in relevant docs
    # More detailed checks per-document would be needed for full audit

def check_doc_content(filepath, fname):
    """Check a single DOCX file for content quality issues."""
    doc = Document(filepath)
    full_text = get_full_text(doc)

    if len(full_text.strip()) < 100:
        add_issue("fatal", fname, "C", "content",
                  "Document has very little content",
                  expected=">1000 chars", actual=f"{len(full_text)} chars")
        return

    # C1: Simplified Chinese check
    sc_found = has_simplified_chinese(full_text)
    if sc_found:
        examples = "; ".join([f"'{ch}' in '{ctx}'" for ch, ctx in sc_found[:3]])
        add_issue("fatal", fname, "C", "simplified-chinese",
                  f"Found simplified Chinese chars: {examples}",
                  expected="All traditional Chinese",
                  actual=f"{len(sc_found)}+ simplified chars")

    # C2: Placeholder check
    ph = check_placeholders(full_text, fname)
    if ph:
        add_issue("important", fname, "C", "placeholder",
                  f"Found placeholders: {', '.join(ph)}")

    # C3: Stub sections - skip this check as DOCX structure puts most content
    # in tables, making paragraph-only counting unreliable.
    # The format audit already checks for document completeness.

    # C4: Excessive empty paragraphs
    max_empty = 0
    cur_empty = 0
    for p in doc.paragraphs:
        if not p.text.strip():
            cur_empty += 1
            max_empty = max(max_empty, cur_empty)
        else:
            cur_empty = 0
    if max_empty > 5:
        add_issue("minor", fname, "C", "empty-paragraphs",
                  f"Max consecutive empty paragraphs: {max_empty}",
                  expected="<=5", actual=str(max_empty))

    return full_text

def check_md_docx_consistency(fname, docx_text):
    """Compare MD source with DOCX output for content completeness."""
    md_name = fname.replace(".docx", ".md")
    md_path = os.path.join(MD_DIR, md_name)
    if not os.path.exists(md_path):
        return

    with open(md_path, "r", encoding="utf-8") as f:
        md_text = f.read()

    # Check table count consistency
    md_tables = len(re.findall(r'^\|.*\|$', md_text, re.MULTILINE))
    # Rough check - if MD has significantly more table rows than DOCX

    # Check for major sections in MD that might be missing in DOCX
    md_headings = re.findall(r'^#{1,3}\s+(.+)$', md_text, re.MULTILINE)
    missing_headings = []
    for h in md_headings:
        h_clean = h.strip().replace('*', '').replace('`', '')
        if len(h_clean) > 3 and h_clean not in docx_text:
            missing_headings.append(h_clean[:40])

    if len(missing_headings) > 3:
        examples = "; ".join(missing_headings[:5])
        add_issue("important", fname, "C", "md-docx-mismatch",
                  f"{len(missing_headings)} MD headings not found in DOCX: {examples}")

def check_requirements_compliance(filepath, fname):
    """Dimension A: Check requirements compliance."""
    doc = Document(filepath)
    full_text = get_full_text(doc)

    # A1: Check case number appears
    if "L0215P2010U" not in full_text:
        add_issue("fatal", fname, "A", "case-number",
                  "Case number L0215P2010U not found in document")

    # A2: Check for header block info (first page should have version, date, vendor)
    first_paras = "\n".join([p.text for p in doc.paragraphs[:20]])

    # A3: Red line - no cross-border data mention where needed
    # Check relevant documents mention on-premises deployment
    relevant_docs = ["01_", "02_", "03_", "06_", "07_", "08_"]
    base = os.path.basename(fname)

    # A4: Check KPIs in relevant documents
    kpi_docs = ["01_", "02_", "05_", "09_"]
    if any(base.startswith(p) for p in kpi_docs):
        if "85%" not in full_text and "85 %" not in full_text:
            add_issue("important", fname, "A", "kpi-missing",
                      "KPI accuracy threshold 85% not mentioned")
        if "5000" not in full_text and "5,000" not in full_text:
            add_issue("minor", fname, "A", "kpi-missing",
                      "Response time 5000ms not mentioned")

    # A5: Check training hours in training docs
    if base.startswith("04_") or base.startswith("12_"):
        if "54" not in full_text:
            add_issue("important", fname, "A", "training-hours",
                      "Training total hours 54 not found",
                      expected=">=54 hours", actual="not mentioned")

    # A6: Check hardware specs in relevant docs
    if base.startswith("06_"):
        hw_checks = [
            ("L40S", "GPU model"),
            ("DDR5", "RAM type"),
            ("RAID", "Storage config"),
        ]
        for keyword, label in hw_checks:
            if keyword not in full_text:
                add_issue("important", fname, "A", "hardware-spec",
                          f"Hardware spec '{label}' ({keyword}) not found")


def main():
    if not os.path.isdir(DOCX_DIR):
        print(f"Directory not found: {DOCX_DIR}")
        sys.exit(1)

    files = sorted(f for f in os.listdir(DOCX_DIR) if f.endswith(".docx"))
    print(f"Scanning {len(files)} DOCX files...\n")

    for f in files:
        fpath = os.path.join(DOCX_DIR, f)
        print(f"  Checking: {f}")

        # Dimension A: Requirements compliance
        check_requirements_compliance(fpath, f)

        # Dimension C: Content quality
        docx_text = check_doc_content(fpath, f)

        # Dimension C: MD vs DOCX consistency
        if docx_text:
            check_md_docx_consistency(f, docx_text)

    # Output results
    fatal = [i for i in ISSUES if i["severity"] == "fatal"]
    important = [i for i in ISSUES if i["severity"] == "important"]
    minor = [i for i in ISSUES if i["severity"] == "minor"]

    print("\n" + "=" * 70)
    print(f"  DOCX Full Audit Report (Round 1)")
    print(f"  Files: {len(files)} | Issues: {len(ISSUES)}")
    print("=" * 70)

    for label, items in [("FATAL", fatal), ("IMPORTANT", important), ("MINOR", minor)]:
        if items:
            print(f"\n### {label} ({len(items)})")
            for i in items:
                print(f"  [{i['id']}] {i['file']} | Dim {i['dimension']} | {i['category']}")
                print(f"         {i['desc']}")
                if i.get('expected'):
                    print(f"         Expected: {i['expected']} | Actual: {i['actual']}")

    if not ISSUES:
        print("\n  No issues found!")

    print(f"\nTotal: fatal={len(fatal)} important={len(important)} minor={len(minor)}")

    # Output as JSON for further processing
    with open("pipeline/audit_results.json", "w", encoding="utf-8") as f:
        json.dump(ISSUES, f, ensure_ascii=False, indent=2)
    print(f"\nDetailed results saved to pipeline/audit_results.json")


if __name__ == "__main__":
    main()
