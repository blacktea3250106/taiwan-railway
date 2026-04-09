#!/usr/bin/env python3
"""
排查 docx 格式問題 — 對照政府文書格式參考規範
"""
import os
import sys
from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.oxml.ns import qn

DOCX_DIR = "pipeline/output_docx"
ISSUES = []

def emu_to_cm(emu):
    if emu is None:
        return None
    return round(emu / 360000, 2)

def emu_to_pt(emu):
    if emu is None:
        return None
    return round(emu / 12700, 1)

def add_issue(severity, file, category, desc):
    ISSUES.append({"severity": severity, "file": file, "category": category, "desc": desc})

def check_file(filepath):
    fname = os.path.basename(filepath)
    doc = Document(filepath)

    # ═══ 1. 頁面設定 ═══
    for i, sec in enumerate(doc.sections):
        w = emu_to_cm(sec.page_width)
        h = emu_to_cm(sec.page_height)
        if w and abs(w - 21.0) > 0.5:
            add_issue("致命", fname, "頁面", f"紙張寬度 {w}cm，應為 21.0cm (A4)")
        if h and abs(h - 29.7) > 0.5:
            add_issue("致命", fname, "頁面", f"紙張高度 {h}cm，應為 29.7cm (A4)")

        top = emu_to_cm(sec.top_margin)
        bot = emu_to_cm(sec.bottom_margin)
        left = emu_to_cm(sec.left_margin)
        right = emu_to_cm(sec.right_margin)

        if top and abs(top - 2.54) > 0.5:
            add_issue("重要", fname, "邊界", f"上邊界 {top}cm，建議 2.54cm")
        if bot and abs(bot - 2.54) > 0.5:
            add_issue("重要", fname, "邊界", f"下邊界 {bot}cm，建議 2.54cm")
        if left and left < 2.5:
            add_issue("重要", fname, "邊界", f"左邊界 {left}cm，建議 ≥3.17cm (含裝訂邊)")
        if right and abs(right - 2.54) > 0.5:
            add_issue("重要", fname, "邊界", f"右邊界 {right}cm，建議 2.54cm")

    # ═══ 2. 頁首頁尾 ═══
    sec = doc.sections[0]
    header = sec.header
    h_text = "".join(p.text for p in header.paragraphs).strip() if header else ""
    if not h_text:
        add_issue("重要", fname, "頁首", "缺少頁首")
    elif "臺灣鐵路" not in h_text and "台鐵" not in h_text and "L0215" not in h_text:
        add_issue("一般", fname, "頁首", f"頁首內容不含專案名稱: [{h_text[:50]}]")

    footer = sec.footer
    f_text = "".join(p.text for p in footer.paragraphs).strip() if footer else ""
    if not f_text and not any(
        run._element.findall(qn("w:fldChar"))
        for p in footer.paragraphs for run in p.runs
    ):
        add_issue("重要", fname, "頁尾", "缺少頁尾/頁碼")

    # ═══ 3. 字型檢查 ═══
    # 檢查 Normal 樣式
    normal = doc.styles["Normal"]
    normal_font = normal.font
    normal_size = emu_to_pt(normal_font.size) if normal_font.size else None

    if normal_size and abs(normal_size - 12) > 1:
        add_issue("重要", fname, "字型", f"內文字級 {normal_size}pt，應為 12pt")

    # 檢查東亞字型
    rPr = normal.element.find(qn("w:rPr"))
    if rPr is not None:
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is not None:
            ea_font = rFonts.get(qn("w:eastAsia"))
            if ea_font and "標楷體" not in ea_font and "DFKai" not in ea_font and "BiauKai" not in ea_font:
                add_issue("重要", fname, "字型", f"中文字型為 {ea_font}，應為標楷體")
        else:
            add_issue("重要", fname, "字型", "Normal 樣式未設定東亞字型 (應為標楷體)")
    else:
        add_issue("一般", fname, "字型", "Normal 樣式缺少 rPr 元素")

    # 檢查英文字型
    if normal_font.name and "Times New Roman" not in normal_font.name:
        add_issue("一般", fname, "字型", f"英文字型為 {normal_font.name}，建議 Times New Roman")

    # ═══ 4. 標題樣式 ═══
    for hname, expected_min, expected_max in [
        ("Heading 1", 16, 22), ("Heading 2", 14, 18), ("Heading 3", 12, 16)
    ]:
        try:
            style = doc.styles[hname]
            sz = emu_to_pt(style.font.size) if style.font.size else None
            if sz and (sz < expected_min or sz > expected_max):
                add_issue("一般", fname, "標題", f"{hname} 字級 {sz}pt，建議 {expected_min}~{expected_max}pt")
            if style.font.bold is not True:
                add_issue("一般", fname, "標題", f"{hname} 未設粗體")
        except KeyError:
            add_issue("一般", fname, "標題", f"缺少 {hname} 樣式")

    # ═══ 5. 表格框線 ═══
    tables = doc.tables
    tables_no_border = 0
    for t in tables:
        tblPr = t._tbl.find(qn("w:tblPr"))
        if tblPr is not None:
            tblBorders = tblPr.find(qn("w:tblBorders"))
            if tblBorders is None:
                tables_no_border += 1
    if tables_no_border > 0:
        add_issue("致命", fname, "表格", f"{tables_no_border}/{len(tables)} 個表格缺少框線")

    # ═══ 6. 目錄(TOC) ═══
    has_toc = False
    for p in doc.paragraphs[:30]:
        if "目錄" in p.text or p.style.name.startswith("toc") or p.style.name.startswith("TOC"):
            has_toc = True
            break
        # 也檢查 field code
        for run in p.runs:
            for elem in run._element.findall(qn("w:instrText")):
                if elem.text and "TOC" in elem.text:
                    has_toc = True
                    break
    if not has_toc:
        add_issue("重要", fname, "目錄", "前 30 段未找到目錄 (TOC)")

    # ═══ 7. 段落行距 ═══
    sample_paras = [p for p in doc.paragraphs if p.text.strip() and p.style.name == "Normal"][:10]
    for p in sample_paras:
        spacing = p.paragraph_format.line_spacing
        if spacing is not None:
            if isinstance(spacing, Pt):
                pass  # fixed spacing, ok
            elif isinstance(spacing, (int, float)):
                if spacing > 3.0:
                    add_issue("一般", fname, "行距", f"行距過大: {spacing}")

    # ═══ 8. 空段落過多 ═══
    consecutive_empty = 0
    max_empty = 0
    for p in doc.paragraphs:
        if not p.text.strip():
            consecutive_empty += 1
            max_empty = max(max_empty, consecutive_empty)
        else:
            consecutive_empty = 0
    if max_empty > 5:
        add_issue("一般", fname, "格式", f"連續空段落最多 {max_empty} 個")

    return len(tables)


def main():
    if not os.path.isdir(DOCX_DIR):
        print(f"目錄不存在: {DOCX_DIR}")
        sys.exit(1)

    files = sorted(f for f in os.listdir(DOCX_DIR) if f.endswith(".docx"))
    print(f"檢測 {len(files)} 個 docx 檔案...\n")

    total_tables = 0
    for f in files:
        n = check_file(os.path.join(DOCX_DIR, f))
        total_tables += n

    # 輸出報告
    fatal = [i for i in ISSUES if i["severity"] == "致命"]
    important = [i for i in ISSUES if i["severity"] == "重要"]
    minor = [i for i in ISSUES if i["severity"] == "一般"]

    print("=" * 60)
    print(f"  格式排查報告")
    print(f"  檔案數: {len(files)} | 表格數: {total_tables}")
    print("=" * 60)

    for label, items in [("致命級", fatal), ("重要級", important), ("一般級", minor)]:
        if items:
            print(f"\n### {label} ({len(items)} 個)")
            for i in items:
                print(f"  [{i['severity']}] {i['file']} | {i['category']} | {i['desc']}")

    print(f"\n總計: 致命={len(fatal)} 重要={len(important)} 一般={len(minor)}")

    if len(fatal) == 0 and len(important) == 0 and len(minor) <= 3:
        print("\n✅ 格式品質達標！")
    else:
        print(f"\n❌ 尚有 {len(fatal)} 致命 + {len(important)} 重要 問題需修復")


if __name__ == "__main__":
    main()
