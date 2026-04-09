#!/usr/bin/env python3
"""
DOCX 深度內容稽核 — 維度 A（需求合規）+ 維度 B（格式補充）+ 維度 C（內容品質）
對照：台鐵需求文件.docx / decisions.yaml / pipeline/output/*.md
"""
import os
import re
import sys
import yaml
from docx import Document
from docx.oxml.ns import qn

DOCX_DIR = "pipeline/output_docx"
MD_DIR = "pipeline/output"
DECISIONS_FILE = "pipeline/decisions.yaml"
REQ_FILE = "台鐵需求文件.docx"

ISSUES = []

# ─── 嚴重度常量 ───
FATAL = "致命"
IMPORTANT = "重要"
MINOR = "一般"

def add_issue(severity, file, dimension, category, desc, expected="", actual=""):
    entry = {
        "severity": severity,
        "file": file,
        "dimension": dimension,
        "category": category,
        "desc": desc,
    }
    if expected:
        entry["expected"] = expected
    if actual:
        entry["actual"] = actual
    ISSUES.append(entry)


def load_decisions():
    with open(DECISIONS_FILE, "r", encoding="utf-8") as f:
        return yaml.safe_load(f)


def load_req_doc():
    """提取需求文件全文文字"""
    doc = Document(REQ_FILE)
    return "\n".join(p.text for p in doc.paragraphs if p.text.strip())


def get_docx_text(filepath):
    """提取 DOCX 全文（段落 + 表格）"""
    doc = Document(filepath)
    parts = []
    for p in doc.paragraphs:
        parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


def get_docx_paragraphs(filepath):
    doc = Document(filepath)
    return [p.text for p in doc.paragraphs]


def get_docx_tables(filepath):
    doc = Document(filepath)
    tables_data = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        tables_data.append(rows)
    return tables_data


def get_md_text(docx_fname):
    """取得對應的 Markdown 原稿文字"""
    md_name = docx_fname.replace(".docx", ".md")
    md_path = os.path.join(MD_DIR, md_name)
    if not os.path.exists(md_path):
        return None
    with open(md_path, "r", encoding="utf-8") as f:
        return f.read()


# ═══════════════════════════════════════════════════════════════
#  簡體中文檢測
# ═══════════════════════════════════════════════════════════════

# 常見簡體字 vs 繁體對照（只列高頻出現的）
SIMPLIFIED_CHARS = set(
    "关于这个来说对与并从进将会为着把给还没被让里面后边头发现"
    "关闭进行运行机器设备数据库处理过程内存调试实现获取无论输"
    "点击响应识别确认记录关联项目数量测试执行报告书体规范"
    "继续创建团队组织结构达标应该规划选择范围环境变量"
)

# 更精準的簡體偵測：只檢查「繁簡不同形」的字
# 這些字在繁體中不會出現（或極罕見），出現即為簡體
DEFINITE_SIMPLIFIED = {
    '关': '關', '对': '對', '这': '這', '来': '來', '进': '進',
    '将': '將', '为': '為', '让': '讓', '从': '從', '还': '還',
    '没': '沒', '发': '發', '现': '現', '头': '頭',
    '边': '邊', '后': '後', '会': '會', '与': '與', '并': '並',
    '给': '給', '着': '著', '过': '過', '获': '獲', '运': '運',
    '设': '設', '备': '備', '数': '數', '处': '處', '内': '內',
    '实': '實', '创': '創', '继': '繼', '续': '續', '选': '選',
    '择': '擇', '规': '規', '达': '達', '应': '應', '环': '環',
    '组': '組', '织': '織', '结': '結', '构': '構', '团': '團',
    '执': '執', '认': '認', '记': '記', '录': '錄', '体': '體',
    '库': '庫', '调': '調', '试': '試', '测': '測', '响': '響',
    '识': '識', '项': '項', '输': '輸', '确': '確', '闭': '閉',
    '击': '擊', '报': '報', '标': '標', '无': '無',
    '论': '論', '点': '點', '范': '範', '画': '畫', '图': '圖',
    '联': '聯', '开': '開', '变': '變',
}
# 繁簡同形字（不應報為簡體）：量、里（里程碑）

def check_simplified_chinese(text, fname):
    """檢查是否包含簡體中文字元"""
    found = {}
    for i, ch in enumerate(text):
        if ch in DEFINITE_SIMPLIFIED:
            # 排除繁簡同形但意義不同的字（需上下文判斷）
            # '里' 在 '里程碑'、'公里'、'千里' 等詞中是繁體正確用法
            if ch == '里':
                ctx = text[max(0, i-2):i+3]
                if any(w in ctx for w in ['里程', '公里', '千里', '鄰里', '里長', '里民']):
                    continue
            trad = DEFINITE_SIMPLIFIED[ch]
            if ch not in found:
                start = max(0, i - 10)
                end = min(len(text), i + 10)
                context = text[start:end].replace('\n', ' ')
                found[ch] = (trad, context)
    return found


# ═══════════════════════════════════════════════════════════════
#  大陸品牌檢測
# ═══════════════════════════════════════════════════════════════

CHINA_BRANDS = [
    "華為", "Huawei", "中興", "ZTE", "海康威視", "Hikvision",
    "大華", "Dahua", "聯想", "Lenovo", "小米", "Xiaomi",
    "百度", "Baidu", "阿里", "Alibaba", "騰訊", "Tencent",
    "科大訊飛", "iFlytek", "商湯", "SenseTime", "曠視", "Megvii",
    "浪潮", "Inspur", "銳捷", "Ruijie", "深信服", "Sangfor",
    "紫光", "H3C", "新華三", "螞蟻", "Ant Design", "antd",
]

CHINA_TERMS = [
    "软件", "硬件", "数据库", "服务器", "网络",
    "视频", "音频", "U盘",
]
# 已排除的詞：
# - '信息' — 可能出現在「排除大陸用語」對照表中
# - '程序' — 繁體中文正常用法（法律程序、變更程序）
# - '接口' — 台灣也常用（API 接口）


# ═══════════════════════════════════════════════════════════════
#  Placeholder 檢測
# ═══════════════════════════════════════════════════════════════

PLACEHOLDER_PATTERNS = [
    r"TODO", r"TBD", r"待補", r"待填", r"xxx", r"XXX",
    r"\[填入\]", r"\[待補\]", r"\[TBD\]", r"【\s*】",
    r"____", r"\.\.\.\.\.",
]


# ═══════════════════════════════════════════════════════════════
#  KPI 數字驗證
# ═══════════════════════════════════════════════════════════════

KPI_CHECKS = {
    "85%": ["辨識率", "準確率", "正確率", "accuracy", "分類準確"],
    "5000": ["回應時間", "回應延遲", "response time", "毫秒", "ms"],
    "30": ["並行", "concurrent", "席次", "port"],
    "99%": ["可用率", "availability", "SLA"],
    "54": ["訓練時數", "教育訓練", "training"],
    "4小時": ["故障修復", "MTTR", "修復時間"],
    "10": ["客製報表", "custom report"],
}


# ═══════════════════════════════════════════════════════════════
#  表格完整性檢測
# ═══════════════════════════════════════════════════════════════

def check_table_integrity(filepath, fname):
    """檢查表格是否有空白儲存格、截斷等問題"""
    doc = Document(filepath)
    for tidx, table in enumerate(doc.tables):
        empty_cells = 0
        total_cells = 0
        for row in table.rows:
            for cell in row.cells:
                total_cells += 1
                if not cell.text.strip():
                    empty_cells += 1
        if total_cells > 0 and empty_cells / total_cells > 0.5:
            add_issue(IMPORTANT, fname, "C", "表格完整性",
                      f"表格 #{tidx+1} 有 {empty_cells}/{total_cells} 個空白儲存格 (>{50}%)")

        # 檢查跨頁表格標題列重複 (透過 tblHeader 屬性)
        first_row = table.rows[0] if table.rows else None
        if first_row and len(table.rows) > 10:
            trPr = first_row._tr.find(qn("w:trPr"))
            has_header = False
            if trPr is not None:
                tblHeader = trPr.find(qn("w:tblHeader"))
                if tblHeader is not None:
                    has_header = True
            if not has_header:
                add_issue(MINOR, fname, "B", "表格格式",
                          f"表格 #{tidx+1} ({len(table.rows)} 列) 未設定標題列重複 (跨頁可能無表頭)")


# ═══════════════════════════════════════════════════════════════
#  Header Block 檢測
# ═══════════════════════════════════════════════════════════════

HEADER_BLOCK_KEYWORDS = ["版本", "案號", "L0215P2010U", "日期"]


# ═══════════════════════════════════════════════════════════════
#  主掃描邏輯
# ═══════════════════════════════════════════════════════════════

def scan_file(filepath, fname, decisions, req_text):
    text = get_docx_text(filepath)
    paragraphs = get_docx_paragraphs(filepath)
    md_text = get_md_text(fname)

    # ─── 維度 A：合規紅線 ───

    # A1. 簡體中文
    simplified = check_simplified_chinese(text, fname)
    if simplified:
        samples = list(simplified.items())[:5]
        detail = "; ".join(f"'{s}' (應為 '{t}') 上下文: ...{ctx}..." for s, (t, ctx) in samples)
        add_issue(FATAL, fname, "A", "簡體中文",
                  f"發現 {len(simplified)} 個簡體字: {detail}")

    # A2. 大陸品牌（排除「排除/禁止/不予採用」語境）
    EXCLUSION_KEYWORDS = ["排除", "禁止", "不予採用", "不得", "違規", "不使用", "因開發組織屬大陸"]
    for brand in CHINA_BRANDS:
        if brand.lower() in text.lower():
            idx = text.lower().find(brand.lower())
            context = text[max(0, idx-80):idx+len(brand)+80].replace('\n', ' ')
            # 如果上下文包含排除語義，跳過
            if any(kw in context for kw in EXCLUSION_KEYWORDS):
                continue
            add_issue(FATAL, fname, "A", "大陸品牌",
                      f"發現大陸品牌 '{brand}'", "不得使用大陸品牌", f"...{context}...")

    # A3. 大陸用語（排除出現在對照表/排除語境中的）
    TERM_EXCLUSION = ["→", "對照", "禁止", "不使用", "避免", "改為"]
    for term in CHINA_TERMS:
        if term in text:
            idx = text.find(term)
            context = text[max(0, idx-30):idx+len(term)+30].replace('\n', ' ')
            if any(kw in context for kw in TERM_EXCLUSION):
                continue
            add_issue(IMPORTANT, fname, "A", "大陸用語",
                      f"疑似大陸用語 '{term}'", "應使用繁體台灣用語", f"...{context}...")

    # A4. Placeholder（去重：同一 pattern 只報一次）
    placeholder_found = set()
    for pattern in PLACEHOLDER_PATTERNS:
        matches = list(re.finditer(pattern, text, re.IGNORECASE))
        if matches:
            key = matches[0].group().upper()
            if key in placeholder_found:
                continue
            placeholder_found.add(key)
            sample = matches[0]
            start = max(0, sample.start() - 20)
            end = min(len(text), sample.end() + 20)
            context = text[start:end].replace('\n', ' ')
            add_issue(IMPORTANT, fname, "A", "Placeholder",
                      f"發現 placeholder '{sample.group()}' ({len(matches)} 處)",
                      "無 placeholder", f"...{context}...")

    # A5. 案號
    if "L0215P2010U" not in text:
        add_issue(IMPORTANT, fname, "A", "案號",
                  "文件內容未包含案號 L0215P2010U")

    # ─── 維度 B：格式補充（audit_docx_format.py 未覆蓋的部分）───

    # B1. Header Block
    first_30_paras = paragraphs[:30]
    first_text = "\n".join(first_30_paras)
    missing_header = [kw for kw in HEADER_BLOCK_KEYWORDS if kw not in first_text]
    if missing_header:
        add_issue(IMPORTANT, fname, "B", "Header Block",
                  f"首頁缺少: {', '.join(missing_header)}",
                  "含文件名稱、版本、案號、日期")

    # B2. 章節編號連續性（只檢查 Heading 樣式的段落）
    doc_obj = Document(filepath)
    heading_nums = []
    for p in doc_obj.paragraphs:
        if p.style and p.style.name and p.style.name.startswith("Heading"):
            m = re.match(r'^(\d+(?:\.\d+)*)\s', p.text.strip())
            if m:
                heading_nums.append(m.group(1))
    # 檢查頂層編號連續性
    top_level = [h for h in heading_nums if '.' not in h]
    for i in range(1, len(top_level)):
        try:
            curr = int(top_level[i])
            prev = int(top_level[i-1])
            if curr != prev + 1 and curr != prev:
                add_issue(MINOR, fname, "B", "編號連續",
                          f"章節編號跳號: {prev} → {curr}")
        except ValueError:
            pass

    # B3. 表格完整性
    check_table_integrity(filepath, fname)

    # ─── 維度 C：內容品質 ───

    # C1. 空白/stub 章節（排除 TOC 區域）
    doc_obj2 = Document(filepath)
    in_toc = False
    past_toc = False
    current_heading = None
    heading_content_lines = 0
    for p in doc_obj2.paragraphs:
        # 偵測 TOC 區域（目錄段落通常使用 toc/TOC/Compact 樣式）
        style_name = p.style.name if p.style else ""
        if style_name.startswith(("toc", "TOC", "Compact")) and not past_toc:
            in_toc = True
            continue
        if in_toc and not style_name.startswith(("toc", "TOC", "Compact")):
            in_toc = False
            past_toc = True
        if in_toc:
            continue

        text = p.text.strip()
        if re.match(r'^#{1,3}\s|^第[一二三四五六七八九十]+章|^\d+\.\s', text):
            if current_heading and heading_content_lines < 3 and heading_content_lines > 0:
                add_issue(IMPORTANT, fname, "C", "Stub 章節",
                          f"章節 '{current_heading[:30]}' 實質內容不足 ({heading_content_lines} 行)")
            current_heading = text[:50]
            heading_content_lines = 0
        elif text:
            heading_content_lines += 1

    # C2. 用語統一
    terms_count = {}
    for term in ["廠商", "本公司", "立約商", "得標商", "承包商"]:
        count = text.count(term)
        if count > 0:
            terms_count[term] = count
    if len(terms_count) > 2:
        detail = ", ".join(f"'{t}'({c}次)" for t, c in terms_count.items())
        add_issue(MINOR, fname, "C", "用語不統一",
                  f"多種廠商稱謂混用: {detail}", "統一使用一種稱謂")

    # C3. Markdown→DOCX 內容遺漏
    if md_text:
        # 精確計算 Markdown 表格數（找 header separator 行 |---|）
        md_table_count = len(re.findall(r'^\|[\s\-:|]+\|$', md_text, re.MULTILINE))

        doc = Document(filepath)
        docx_table_count = len(doc.tables)

        # 差異超過 30% 才報（Pandoc 可能合併相鄰表格）
        if md_table_count > 0 and docx_table_count < md_table_count * 0.5:
            add_issue(IMPORTANT, fname, "C", "轉換遺漏",
                      f"Markdown {md_table_count} 個表格，DOCX 只有 {docx_table_count} 個",
                      f"{md_table_count} 個表格", f"{docx_table_count} 個表格")

    # C4. 連續空段落
    consecutive = 0
    max_consecutive = 0
    for p in paragraphs:
        if not p.strip():
            consecutive += 1
            max_consecutive = max(max_consecutive, consecutive)
        else:
            consecutive = 0
    if max_consecutive > 5:
        add_issue(MINOR, fname, "C", "空段落",
                  f"連續空段落最多 {max_consecutive} 個", "≤5", str(max_consecutive))


def scan_cross_file_kpi(decisions):
    """跨文件 KPI 數字一致性"""
    kpi = decisions.get("kpi", {})
    files = sorted(f for f in os.listdir(DOCX_DIR) if f.endswith(".docx"))

    # 特定 KPI 在特定文件中必須出現
    kpi_file_map = {
        "05_測試計畫書.docx": ["85%", "5000", "99%"],
        "09_測試報告書.docx": ["85%", "5000", "99%"],
        "01_工作計畫書.docx": ["54"],
        "04_教育訓練計畫書.docx": ["54"],
        "13_結案報告.docx": ["85%", "99%"],
    }

    for fname, expected_kpis in kpi_file_map.items():
        fpath = os.path.join(DOCX_DIR, fname)
        if not os.path.exists(fpath):
            continue
        text = get_docx_text(fpath)
        for kpi_val in expected_kpis:
            # 支援多種數字格式（5000 / 5,000 / 5000ms）
            variants = [kpi_val]
            if kpi_val == "5000":
                variants = ["5000", "5,000"]
            if kpi_val == "54":
                variants = ["54"]
            found = any(v in text for v in variants)
            if not found:
                kpi_desc = {
                    "85%": "辨識率/準確率 ≥85%",
                    "5000": "回應時間 ≤5000ms",
                    "99%": "可用率 ≥99%",
                    "54": "訓練時數 ≥54小時",
                }.get(kpi_val, kpi_val)
                add_issue(IMPORTANT, fname, "A", "KPI 缺失",
                          f"未找到 KPI 指標 '{kpi_desc}'",
                          f"應包含 {kpi_val}", "未出現")


def scan_hardware_consistency(decisions):
    """硬體規格跨文件一致性"""
    hw = decisions.get("hardware", {}).get("production", {})
    hw_files = ["06_軟硬體清單及系統架構圖.docx", "02_系統分析報告書.docx"]

    for fname in hw_files:
        fpath = os.path.join(DOCX_DIR, fname)
        if not os.path.exists(fpath):
            continue
        text = get_docx_text(fpath)

        # 檢查 AI 伺服器 GPU
        if "L40S" not in text and "AI" in fname or "系統分析" in fname:
            # L40S 在 AI 相關文件中必須出現
            if "L40S" not in text:
                add_issue(IMPORTANT, fname, "A", "硬體規格",
                          "未提及 GPU 型號 NVIDIA L40S",
                          "NVIDIA L40S 48GB", "未出現")

        # AI server 數量
        ai_qty = hw.get("ai_server", {}).get("qty", 2)
        if f"{ai_qty}" not in text and "AI 伺服器" in text:
            pass  # 數字可能以其他形式呈現，不誤報


def main():
    if not os.path.isdir(DOCX_DIR):
        print(f"目錄不存在: {DOCX_DIR}")
        sys.exit(1)

    print("載入 decisions.yaml...")
    decisions = load_decisions()

    print("載入台鐵需求文件...")
    req_text = load_req_doc()

    files = sorted(f for f in os.listdir(DOCX_DIR) if f.endswith(".docx"))
    print(f"\n掃描 {len(files)} 個 DOCX 檔案...\n")

    for fname in files:
        fpath = os.path.join(DOCX_DIR, fname)
        print(f"  掃描: {fname}")
        scan_file(fpath, fname, decisions, req_text)

    print("\n跨文件 KPI 一致性檢查...")
    scan_cross_file_kpi(decisions)

    print("硬體規格一致性檢查...")
    scan_hardware_consistency(decisions)

    # ── 輸出報告 ──
    fatal = [i for i in ISSUES if i["severity"] == FATAL]
    important = [i for i in ISSUES if i["severity"] == IMPORTANT]
    minor = [i for i in ISSUES if i["severity"] == MINOR]

    print("\n" + "=" * 60)
    print("  DOCX 深度內容稽核報告")
    print("=" * 60)

    issue_id = 0
    for label, items in [("致命級", fatal), ("重要級", important), ("一般級", minor)]:
        if items:
            print(f"\n### {label} ({len(items)} 個)")
            for i in items:
                issue_id += 1
                line = f"  [DOCX-{issue_id:03d}] {i['file']} | 維度{i['dimension']} | {i['category']} | {i['desc']}"
                if i.get("expected"):
                    line += f"\n           預期: {i['expected']}"
                if i.get("actual"):
                    line += f"\n           實際: {i['actual']}"
                print(line)

    print(f"\n總計: 致命={len(fatal)} 重要={len(important)} 一般={len(minor)}")

    if len(fatal) == 0 and len(important) == 0 and len(minor) <= 3:
        print("\n✅ 內容品質達標！")
    else:
        print(f"\n❌ 尚有 {len(fatal)} 致命 + {len(important)} 重要 問題需修復")

    return len(fatal), len(important), len(minor)


if __name__ == "__main__":
    main()
