#!/usr/bin/env python3
"""
建立政府標案交付文件 reference.docx 模板
依據：國發會《政府文書格式參考規範》

用途：作為 pandoc --reference-doc 的模板
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy


def set_run_font(run, cn_font="標楷體", en_font="Times New Roman", size=12):
    """設定 run 的中英文字型與大小"""
    run.font.size = Pt(size)
    run.font.name = en_font
    # 設定東亞字型（中文）
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), cn_font)
    rFonts.set(qn("w:ascii"), en_font)
    rFonts.set(qn("w:hAnsi"), en_font)


def set_style_font(style, cn_font="標楷體", en_font="Times New Roman", size=12,
                   bold=False, color=None, space_before=0, space_after=0,
                   line_spacing=None, alignment=None, keep_next=False):
    """設定段落樣式的完整屬性"""
    fmt = style.paragraph_format
    fmt.space_before = Pt(space_before)
    fmt.space_after = Pt(space_after)
    if line_spacing:
        fmt.line_spacing = Pt(line_spacing)
    if alignment is not None:
        fmt.alignment = alignment
    fmt.keep_with_next = keep_next

    # 字型設定
    font = style.font
    font.name = en_font
    font.size = Pt(size)
    font.bold = bold
    if color:
        font.color.rgb = RGBColor(*color)

    # 東亞字型
    rPr = style.element.find(qn("w:rPr"))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")} />')
        style.element.append(rPr)
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} />')
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), cn_font)
    rFonts.set(qn("w:ascii"), en_font)
    rFonts.set(qn("w:hAnsi"), en_font)


def set_page_margins(section, top=2.54, bottom=2.54, left=3.17, right=2.54):
    """設定頁面邊界（cm）— 左邊含裝訂邊"""
    section.top_margin = Cm(top)
    section.bottom_margin = Cm(bottom)
    section.left_margin = Cm(left)
    section.right_margin = Cm(right)


def add_page_number(paragraph):
    """在段落中插入頁碼欄位"""
    run = paragraph.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._element.append(fldChar1)

    run2 = paragraph.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._element.append(instrText)

    run3 = paragraph.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._element.append(fldChar2)


def add_total_pages(paragraph):
    """在段落中插入總頁數欄位"""
    run = paragraph.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._element.append(fldChar1)

    run2 = paragraph.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> NUMPAGES </w:instrText>')
    run2._element.append(instrText)

    run3 = paragraph.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._element.append(fldChar2)


def setup_header_footer(section, project_name="臺灣鐵路管理局 AI 客服系統建置案（案號 L0215P2010U）"):
    """設定頁首頁尾"""
    # ── 頁首 ──
    header = section.header
    header.is_linked_to_previous = False
    if header.paragraphs:
        hp = header.paragraphs[0]
    else:
        hp = header.add_paragraph()
    hp.clear()
    run = hp.add_run(project_name)
    set_run_font(run, size=9)
    run.font.color.rgb = RGBColor(128, 128, 128)
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 頁首底線
    pPr = hp._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="999999"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

    # ── 頁尾 ──
    footer = section.footer
    footer.is_linked_to_previous = False
    if footer.paragraphs:
        fp = footer.paragraphs[0]
    else:
        fp = footer.add_paragraph()
    fp.clear()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run_prefix = fp.add_run("第 ")
    set_run_font(run_prefix, size=9)
    run_prefix.font.color.rgb = RGBColor(128, 128, 128)

    add_page_number(fp)

    run_mid = fp.add_run(" 頁，共 ")
    set_run_font(run_mid, size=9)
    run_mid.font.color.rgb = RGBColor(128, 128, 128)

    add_total_pages(fp)

    run_suffix = fp.add_run(" 頁")
    set_run_font(run_suffix, size=9)
    run_suffix.font.color.rgb = RGBColor(128, 128, 128)


def setup_table_style(doc):
    """建立表格樣式（有框線）"""
    # Pandoc 使用 "Table" 或 "Table Grid" 樣式
    # 修改預設的 Table Grid
    for style in doc.styles:
        if style.name == "Table Grid":
            return  # 已存在就不重複建立

    # 如果 Table Grid 不存在，使用已有的 Table Normal
    pass


def create_reference():
    doc = Document()

    # ══════════════════════════════════════
    #  頁面設定
    # ══════════════════════════════════════
    section = doc.sections[0]
    section.page_width = Cm(21.0)   # A4
    section.page_height = Cm(29.7)  # A4
    set_page_margins(section, top=2.54, bottom=2.54, left=3.17, right=2.54)

    # ══════════════════════════════════════
    #  樣式設定
    # ══════════════════════════════════════

    # Normal（內文）：標楷體 12pt，行距約 18pt（1.5 倍單行）
    set_style_font(
        doc.styles["Normal"],
        size=12, line_spacing=18,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    )

    # Heading 1：第一章等級，20pt 粗體
    set_style_font(
        doc.styles["Heading 1"],
        size=20, bold=True,
        space_before=24, space_after=12,
        line_spacing=36, keep_next=True,
    )

    # Heading 2：1.1 等級，16pt 粗體
    set_style_font(
        doc.styles["Heading 2"],
        size=16, bold=True,
        space_before=18, space_after=6,
        line_spacing=28, keep_next=True,
    )

    # Heading 3：1.1.1 等級，14pt 粗體
    set_style_font(
        doc.styles["Heading 3"],
        size=14, bold=True,
        space_before=12, space_after=6,
        line_spacing=24, keep_next=True,
    )

    # Heading 4
    set_style_font(
        doc.styles["Heading 4"],
        size=13, bold=True,
        space_before=6, space_after=3,
        line_spacing=20, keep_next=True,
    )

    # Heading 5
    set_style_font(
        doc.styles["Heading 5"],
        size=12, bold=True,
        space_before=6, space_after=3,
        line_spacing=18, keep_next=True,
    )

    # Body Text
    if "Body Text" in [s.name for s in doc.styles]:
        set_style_font(doc.styles["Body Text"], size=12, line_spacing=18)

    # First Paragraph（Pandoc 用這個）
    try:
        set_style_font(doc.styles["First Paragraph"], size=12, line_spacing=18)
    except KeyError:
        pass

    # Compact（Pandoc list style）
    try:
        compact = doc.styles.add_style("Compact", 1)  # paragraph style
        set_style_font(compact, size=12, line_spacing=16,
                       space_before=2, space_after=2)
    except ValueError:
        pass

    # TOC 樣式
    for level in range(1, 4):
        toc_name = f"toc {level}"
        try:
            style = doc.styles[toc_name]
            set_style_font(style, size=12 - (level - 1),
                           space_before=2, space_after=2)
        except KeyError:
            pass

    # ══════════════════════════════════════
    #  頁首頁尾
    # ══════════════════════════════════════
    setup_header_footer(section)

    # ══════════════════════════════════════
    #  佔位內容（Pandoc 需要至少一段文字才能繼承樣式）
    # ══════════════════════════════════════
    p = doc.add_paragraph("")
    # 清除佔位段落（Pandoc 會替換全部內容）

    # ══════════════════════════════════════
    #  儲存
    # ══════════════════════════════════════
    output_path = "pipeline/reference.docx"
    doc.save(output_path)
    print(f"✅ reference.docx 建立完成: {output_path}")
    print(f"   字型: 標楷體 / Times New Roman")
    print(f"   內文: 12pt, 行距 18pt")
    print(f"   標題: H1=20pt, H2=16pt, H3=14pt")
    print(f"   邊界: 上下 2.54cm, 左 3.17cm (含裝訂邊), 右 2.54cm")
    print(f"   頁首: 專案名稱")
    print(f"   頁尾: 第 X 頁，共 Y 頁")


if __name__ == "__main__":
    create_reference()
