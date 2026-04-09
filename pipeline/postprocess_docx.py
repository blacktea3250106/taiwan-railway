#!/usr/bin/env python3
"""
DOCX 後處理腳本 — 修復 Pandoc 轉換後的格式問題
1. 為所有表格加上框線
2. 插入 Word 原生 TOC 欄位
3. 設定表格內文字型為標楷體 10pt
"""
import os
import sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Emu
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_table_borders(table):
    """為表格加上完整框線（單線、黑色、0.5pt）"""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)

    # 移除既有 borders
    existing = tblPr.find(qn("w:tblBorders"))
    if existing is not None:
        tblPr.remove(existing)

    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)


def set_table_font(table, cn_font="標楷體", en_font="Times New Roman", size=10):
    """設定表格內所有文字的字型"""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(size)
                    run.font.name = en_font
                    rPr = run._element.get_or_add_rPr()
                    rFonts = rPr.find(qn("w:rFonts"))
                    if rFonts is None:
                        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")}/>')
                        rPr.insert(0, rFonts)
                    rFonts.set(qn("w:eastAsia"), cn_font)
                    rFonts.set(qn("w:ascii"), en_font)
                    rFonts.set(qn("w:hAnsi"), en_font)


def shade_header_row(table):
    """為表格第一列加上淺灰底色 + 粗體"""
    if len(table.rows) == 0:
        return
    for cell in table.rows[0].cells:
        # 底色
        tcPr = cell._element.find(qn("w:tcPr"))
        if tcPr is None:
            tcPr = parse_xml(f'<w:tcPr {nsdecls("w")}/>')
            cell._element.insert(0, tcPr)
        shd = tcPr.find(qn("w:shd"))
        if shd is None:
            shd = parse_xml(
                f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="D9E2F3"/>'
            )
            tcPr.append(shd)
        # 粗體
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True


def insert_toc_field(doc):
    """在文件開頭插入 Word 原生 TOC 欄位（需開啟 Word 後按 F9 更新）"""
    # 找到第一個非空段落的位置
    body = doc.element.body
    first_para = body.find(qn("w:p"))
    if first_para is None:
        return

    # 建立「目錄」標題段落
    toc_title = parse_xml(
        f'<w:p {nsdecls("w")}>'
        f'  <w:pPr>'
        f'    <w:jc w:val="center"/>'
        f'    <w:spacing w:after="200"/>'
        f'  </w:pPr>'
        f'  <w:r>'
        f'    <w:rPr>'
        f'      <w:rFonts w:eastAsia="標楷體" w:ascii="Times New Roman" w:hAnsi="Times New Roman"/>'
        f'      <w:b/>'
        f'      <w:sz w:val="32"/>'
        f'    </w:rPr>'
        f'    <w:t>目　錄</w:t>'
        f'  </w:r>'
        f'</w:p>'
    )
    body.insert(list(body).index(first_para), toc_title)

    # 建立 TOC field
    toc_para = parse_xml(
        f'<w:p {nsdecls("w")}>'
        f'  <w:r>'
        f'    <w:fldChar w:fldCharType="begin"/>'
        f'  </w:r>'
        f'  <w:r>'
        f'    <w:instrText xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>'
        f'  </w:r>'
        f'  <w:r>'
        f'    <w:fldChar w:fldCharType="separate"/>'
        f'  </w:r>'
        f'  <w:r>'
        f'    <w:rPr>'
        f'      <w:rFonts w:eastAsia="標楷體"/>'
        f'    </w:rPr>'
        f'    <w:t>（請在 Word 中按 Ctrl+A → F9 更新目錄）</w:t>'
        f'  </w:r>'
        f'  <w:r>'
        f'    <w:fldChar w:fldCharType="end"/>'
        f'  </w:r>'
        f'</w:p>'
    )
    body.insert(list(body).index(first_para), toc_para)

    # TOC 後加分頁符
    page_break = parse_xml(
        f'<w:p {nsdecls("w")}>'
        f'  <w:r>'
        f'    <w:br w:type="page"/>'
        f'  </w:r>'
        f'</w:p>'
    )
    body.insert(list(body).index(first_para), page_break)


def set_header_row_repeat(table):
    """設定表格首行標題列重複（跨頁時自動重複標題列）"""
    rows = table._tbl.findall(qn("w:tr"))
    if len(rows) <= 1:
        return
    first_tr = rows[0]
    trPr = first_tr.find(qn("w:trPr"))
    if trPr is None:
        trPr = parse_xml(f'<w:trPr {nsdecls("w")} />')
        first_tr.insert(0, trPr)
    existing = trPr.find(qn("w:tblHeader"))
    if existing is None:
        tblHeader = parse_xml(f'<w:tblHeader {nsdecls("w")} />')
        trPr.append(tblHeader)


def set_table_width_full(table):
    """設定表格寬度為 100% 頁面寬度"""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    # 設定 auto layout
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="5000" w:type="pct"/>')
        tblPr.append(tblW)
    else:
        tblW.set(qn("w:w"), "5000")
        tblW.set(qn("w:type"), "pct")


def process_file(filepath):
    """處理單一 docx 檔案"""
    fname = os.path.basename(filepath)
    doc = Document(filepath)

    table_count = len(doc.tables)
    changes = []

    # 1. 表格處理
    for table in doc.tables:
        add_table_borders(table)
        set_table_font(table)
        shade_header_row(table)
        set_table_width_full(table)
        set_header_row_repeat(table)
    if table_count > 0:
        changes.append(f"表格框線+字型+標題列底色 ({table_count} 個)")

    # 2. TOC 插入（檢查是否已有）
    has_toc = False
    for p in doc.paragraphs[:30]:
        if "目錄" in p.text or "目　錄" in p.text:
            has_toc = True
            break
        for run in p.runs:
            for elem in run._element.findall(qn("w:instrText")):
                if elem.text and "TOC" in elem.text:
                    has_toc = True
                    break

    if not has_toc:
        insert_toc_field(doc)
        changes.append("插入 TOC 目錄")

    # 儲存
    doc.save(filepath)
    return changes


def main():
    docx_dir = sys.argv[1] if len(sys.argv) > 1 else "pipeline/output_docx"

    if not os.path.isdir(docx_dir):
        print(f"目錄不存在: {docx_dir}")
        sys.exit(1)

    files = sorted(f for f in os.listdir(docx_dir) if f.endswith(".docx"))
    print(f"\n後處理 {len(files)} 個 docx 檔案...\n")

    for f in files:
        path = os.path.join(docx_dir, f)
        changes = process_file(path)
        if changes:
            print(f"  ✅ {f}: {', '.join(changes)}")
        else:
            print(f"  ⏭️  {f}: 無需變更")

    print(f"\n後處理完成！")
    print(f"提示: 開啟 Word 後按 Ctrl+A → F9 (Mac: Cmd+A → Fn+F9) 更新目錄頁碼")


if __name__ == "__main__":
    main()
